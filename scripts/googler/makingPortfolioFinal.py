import os
import shutil
from docx import Document
from docx.shared import Pt
from itertools import product
from googleapiclient.http import MediaFileUpload

from Google import folder_create, PORTFOLIO_FOLDER_ID, PORTFOLIO_FORM_SHEET_ID, PORTFOLIO_FORM_SHEET_LIST_NAME, \
    PORTFOLIO_DOC_TEMPLATE_ID, GOOGLE_SHEETS_SERVICE, GOOGLE_DOCS_SERVICE, GOOGLE_DRIVE_SERVICE, give_user_permission

flag_index = 26
flag_symbol = 'Z'

PORTFOLIO_TEMPLATE_FILE_PATH = "/scripts/googler/portfolio_making/portfolio_template_2024.docx"


def get_google_sheets_data():
    # what is range?
    result = GOOGLE_SHEETS_SERVICE.spreadsheets().values().get(
        spreadsheetId=PORTFOLIO_FORM_SHEET_ID, range=PORTFOLIO_FORM_SHEET_LIST_NAME).execute()
    values = result.get('values', [])
    return values


def update_flag_in_sheet(row_index):
    range_name = f'{PORTFOLIO_FORM_SHEET_LIST_NAME}!{flag_symbol}{row_index}'
    # print(range_name)
    value_range_body = {'values': [[str(True)]]}  # Обновляем значение флага
    GOOGLE_SHEETS_SERVICE.spreadsheets().values().update(
        spreadsheetId=PORTFOLIO_FORM_SHEET_ID, range=range_name,
        valueInputOption='RAW', body=value_range_body).execute()


def copy_docx_to_same_folder(original_file_path, copy_name):
    try:
        folder_path, file_name = os.path.split(original_file_path)
        copy_file_path = os.path.join(folder_path, f"{copy_name}.docx")

        shutil.copyfile(original_file_path, copy_file_path)
        return copy_file_path
    except Exception as e:
        print(f"Error copying the file: {e}")
        return None


def upload_docx_to_google_drive(file_path, folder_id):

    folder_path, file_name = os.path.split(file_path)
    file_metadata = {
        'name': os.path.basename(file_name),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = GOOGLE_DRIVE_SERVICE.files().create(body=file_metadata, media_body=media, fields='id').execute()


def make_portfolio(data, folder_id, applicant_name):
    # Cоздаем копию шаблона локально:
    portfolio_name = 'Портфолио ' + applicant_name
    new_portfolio_file_path = copy_docx_to_same_folder(PORTFOLIO_TEMPLATE_FILE_PATH, portfolio_name)

    # Заменяем ключи в копии шаблона:
    # Замена в параграфах
    doc = Document(new_portfolio_file_path)
    for paragraph in doc.paragraphs:
        for search_text, replace_text in data.items():
            if ("#" + search_text + "#") in paragraph.text:
                paragraph.text = paragraph.text.replace("#" + search_text + "#", replace_text)
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for search_text, replace_text in data.items():
                        if ("#" + search_text + "#") in paragraph.text:
                            paragraph.text = paragraph.text.replace("#" + search_text + "#", replace_text)
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)

    doc.save(new_portfolio_file_path)

    # Загружаем итоговую копию в папку на гугл диске:
    f = upload_docx_to_google_drive(new_portfolio_file_path, folder_id)
    # Удаляем локальную копию:
    os.remove(new_portfolio_file_path)
    return f



def main():
    values = get_google_sheets_data()
    for row in values:
        if row == values[0]:  # Пропускаем первую строку с заголовками
            continue
        if len(row) != 0 and len(row) < flag_index: # Выполняем действия только для строк, где флаг пустой
            data = dict(zip(values[0], row))
            applicant_name = data.get('Фамилия') + ' ' + data.get('Имя') + ' ' + data.get('Отчество')

            folder_id = folder_create(applicant_name, PORTFOLIO_FOLDER_ID)
            
            make_portfolio(data, folder_id, applicant_name)

            # Получаем индекс строки (нумерация с 1):
            row_index = values.index(row) + 1
            # Обновляем значение флага в текущей строке:
            update_flag_in_sheet(row_index)
            give_user_permission(folder_id, data.get('Электронная почта'))

if __name__ == '__main__':
    main()
